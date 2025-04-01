import argparse
import os
import sys
import subprocess
import yaml

from docx import Document
from docx.shared import Inches

from dotenv import load_dotenv
from datetime import datetime

# AI Provider Libraries
from openai import OpenAI, OpenAIError
from tavily import TavilyClient
from anthropic import Anthropic, AnthropicError
import google.generativeai as genai
from google.api_core import exceptions as google_exceptions

# Load environment variables from .env file
load_dotenv()

# --- Configuration Loading ---
def load_config(config_path='config.yaml'):
    """Loads the configuration file."""
    try:
        with open(config_path, 'r') as f:
            config = yaml.safe_load(f)
        if not config:
            print(f"Error: Config file '{config_path}' is empty or invalid.")
            sys.exit(1)
        return config
    except FileNotFoundError:
        print(f"Error: Configuration file '{config_path}' not found.")
        sys.exit(1)
    except yaml.YAMLError as e:
        print(f"Error parsing configuration file '{config_path}': {e}")
        sys.exit(1)
    except Exception as e:
        print(f"An unexpected error occurred loading config: {e}")
        sys.exit(1)

# --- API Key Loading ---
def load_api_keys():
    """Loads all required API keys from environment variables."""
    keys = {
        "tavily": os.getenv("TAVILY_API_KEY"),
        "openai": os.getenv("OPENAI_API_KEY"),
        "anthropic": os.getenv("ANTHROPIC_API_KEY"),
        "google": os.getenv("GOOGLE_API_KEY")
    }
    # Basic check - could be more specific based on selected provider later
    if not keys["tavily"]:
        print("Error: TAVILY_API_KEY not found in environment/.env file.")
        sys.exit(1)
    # Specific checks for AI keys will happen during client initialization
    return keys

# --- AI Client Initialization ---
def initialize_ai_client(provider, api_key):
    """Initializes and returns the appropriate AI client based on the provider."""
    if not api_key:
        print(f"Error: API key for provider '{provider}' not found in environment/.env file.")
        sys.exit(1)

    try:
        if provider == 'openai':
            # OpenAI client automatically uses OPENAI_API_KEY env var if not passed explicitly
            return OpenAI(api_key=api_key)
        elif provider == 'anthropic':
            return Anthropic(api_key=api_key)
        elif provider == 'google':
            genai.configure(api_key=api_key)
            # Return the generative model instance configured for the specific model later
            # For now, just return the configured module or a base client if available
            # Let's return the configured module itself for access to GenerativeModel
            return genai
        else:
            print(f"Error: Unsupported AI provider '{provider}' specified in config.")
            sys.exit(1)
    except Exception as e:
        print(f"Error initializing AI client for provider '{provider}': {e}")
        sys.exit(1)

# --- Tavily Client Initialization ---
def initialize_tavily_client(api_key):
    """Initializes the Tavily client."""
    if not api_key:
         print("Error: TAVILY_API_KEY not found for Tavily client.")
         sys.exit(1)
    try:
        return TavilyClient(api_key=api_key)
    except Exception as e:
        print(f"Error initializing Tavily client: {e}")
        sys.exit(1)

# --- Core Functions ---

def search_news(tavily_client, query):
    """
    Searches for news articles using the Tavily client with advanced settings
    and requests associated image URLs.

    Args:
        tavily_client (TavilyClient): An initialized Tavily client instance.
        query (str): The search term or phrase.

    Returns:
        list: A list of dictionaries, where each dictionary represents a
              search result, potentially including an 'images' list.
              Returns an empty list on failure.
    """
    try:
        print(f"Performing advanced Tavily search for: '{query}'")
        # Use advanced search, request more results, and include images
        response = tavily_client.search(
            query=query,
            search_depth="advanced", # Use "advanced" for more thorough results
            include_images=True,     # Request associated image URLs
            max_results=10           # Request up to 10 results
            # Optional: Add include_domains or exclude_domains if needed
            # include_domains=["example.com"],
            # exclude_domains=["undesired.com"]
        )
        # Results structure includes 'query', 'answer', 'images', 'results', 'response_time'
        # We primarily need the 'results' list here.
        results = response.get("results", [])
        print(f"Tavily search returned {len(results)} results.")
        # Optional: Log the overall images found at the query level if needed
        # query_level_images = response.get("images", [])
        # if query_level_images:
        #    print(f"Found {len(query_level_images)} query-level images.")
        return results
    except Exception as e:
        print(f"Error during enhanced Tavily search for query '{query}': {e}")
        return []


def format_news_with_ai(ai_client, provider, model_name, results, level):
    """
    Generates a structured news report using the selected AI provider,
    including image URLs found by Tavily in the context provided to the AI,
    and requesting specific delimiters for parsing into DOCX elements.
    """
    news_text_context = "" # Renamed to clarify it's context for the AI
    print("Processing search results for AI context...")

    for i, r in enumerate(results):
        title = r.get("title", "No Title")
        content = r.get("content", "No Content Available")
        url = r.get("url", "No URL Available")
        # Extract the first image URL if available in the result's 'images' list
        image_url = None
        images = r.get("images", []) # Tavily returns 'images' list per result
        if images and isinstance(images, list) and len(images) > 0:
            image_url = images[0]
            print(f"  Result {i+1}: Found image URL: {image_url}")

        # Build the text context for this specific result
        news_text_context += f"--- Source Result {i+1} ---\n"
        news_text_context += f"Source Title: {title}\n"
        news_text_context += f"Source URL: {url}\n"
        if image_url:
            # Include the image URL in the context for the AI
            news_text_context += f"Relevant Image URL: {image_url}\n"
        news_text_context += f"Source Content Snippet: {content}\n\n"


    if not news_text_context:
        print("Warning: No text context could be built from search results.")
        return "Report generation failed: No context from search results."

    # --- Updated Prompt Acknowledging Image URLs ---
    h1_delim = "@@H1@@"
    h2_delim = "@@H2@@"
    p_delim = "@@P@@"

    user_prompt = (
        f"Act as a professional news reporter creating a structured document. "
        f"Based *only* on the following search results (which include text content and potentially a relevant image URL per source), "
        f"create a concise news report for {level} news.\n"
        f"IMPORTANT: Format your entire response using ONLY the following delimiters:\n"
        f"- Use '{h1_delim}' before AND after the single main title of the report.\n"
        f"- Use '{h2_delim}' before AND after each section headline.\n"
        f"- Use '{p_delim}' before AND after each paragraph of text.\n"
        f"Do NOT use any other formatting like Markdown (*, #, -, etc.). "
        f"Mentioning key information from the sources is crucial. Acknowledge images if relevant contextually.\n\n"
        f"Search Results Context:\n===\n{news_text_context}\n===\n"
        f"Begin the structured report now:"
    )

    system_prompt = f"You are an AI assistant that generates structured text documents using specific delimiters: {h1_delim}Title{h1_delim}, {h2_delim}Headline{h2_delim}, {p_delim}Paragraph{p_delim}. You synthesize information from provided text snippets and associated image URLs."

    # --- API Call Logic (remains largely the same, just uses updated prompts) ---
    try:
        print(f"Sending request to {provider.upper()} model {model_name}...")
        if provider == 'openai':
            completion = ai_client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ]
            )
            report = completion.choices[0].message.content
            return report

        elif provider == 'anthropic':
            message = ai_client.messages.create(
                model=model_name,
                max_tokens=4000,
                system=system_prompt,
                messages=[ {"role": "user", "content": user_prompt} ]
            )
            report = message.content[0].text
            return report

        elif provider == 'google':
            if hasattr(ai_client, 'GenerativeModel'):
                 model = ai_client.GenerativeModel(model_name, system_instruction=system_prompt)
            else:
                 print("Error: Google AI client not initialized correctly.")
                 return f"Report generation failed for {level} news (Google Client Init Error)."

            # Relax safety settings slightly for news content - adjust as needed
            safety_settings = [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            ]
            response = model.generate_content(user_prompt, safety_settings=safety_settings)

            if not response.candidates or not response.candidates[0].content.parts:
                 print(f"Warning: Google Gemini response might be empty or blocked for {level} news.")
                 try: print(f"Prompt Feedback: {response.prompt_feedback}")
                 except Exception: pass
                 # Check finish reason if available
                 try:
                     finish_reason = response.candidates[0].finish_reason
                     print(f"Finish Reason: {finish_reason}")
                     if finish_reason != 1: # 1 typically means "STOP"
                         print(f"Content may be incomplete due to finish reason: {finish_reason}")
                 except Exception: pass
                 return f"Report generation failed for {level} news (Google Gemini Blocked/Empty/Safety)."

            report = response.text
            return report
        else:
            return f"Report generation failed: Unsupported provider '{provider}'."

    # --- Error Handling (remains the same) ---
    except OpenAIError as e:
        print(f"OpenAI API error for {level} news: {e}")
        return f"Report generation failed for {level} news (OpenAI Error)."
    except AnthropicError as e:
        print(f"Anthropic API error for {level} news: {e}")
        return f"Report generation failed for {level} news (Anthropic Error)."
    except (google_exceptions.GoogleAPIError, ValueError) as e:
         print(f"Google Gemini API error for {level} news: {e}")
         return f"Report generation failed for {level} news (Google Gemini Error)."
    except Exception as e:
        print(f"Unexpected error generating report with {provider} for {level} news: {e}")
        return f"Report generation failed for {level} news (Unexpected Error)."



def generate_docx(report_text, filename):
    """
    Generates a .docx file by parsing the AI's structured output
    with delimiters (@@H1@@, @@H2@@, @@P@@).

    Args:
        report_text (str): The raw string output from the AI, expected
                           to contain delimiters.
        filename (str): The desired name/path for the output .docx file.

    Returns:
        bool: True if the document was saved successfully, False otherwise.
    """
    # Define the delimiters expected from the AI
    delimiters = {
        "H1": "@@H1@@",
        "H2": "@@H2@@",
        "P": "@@P@@",
        # Add more if needed, e.g., "L": "@@L@@" for lists
    }
    all_delimiters_flat = list(delimiters.values())

    try:
        # Create a new document
        # Optionally load a template: doc = Document('template.docx')
        doc = Document()

        # --- Optional: Set Margins ---
        # sections = doc.sections
        # for section in sections:
        #     section.top_margin = Inches(0.5)
        #     section.bottom_margin = Inches(0.5)
        #     section.left_margin = Inches(0.75)
        #     section.right_margin = Inches(0.75)
        # -----------------------------

        # Split the report text by *any* of the delimiters.
        # We need a way to split while keeping the delimiters or knowing what was split.
        # A simple approach: replace delimiters with unique separators, split, then process.
        # Using a placeholder unlikely to be in the text.
        placeholder = "||SPLIT||"
        processed_text = report_text
        for delim in all_delimiters_flat:
            processed_text = processed_text.replace(delim, f"{placeholder}{delim}{placeholder}")

        parts = [part.strip() for part in processed_text.split(placeholder) if part.strip()]

        if not parts:
             print(f"Warning: No content found after processing delimiters for {filename}.")
             # Optionally save a blank or error document
             doc.add_paragraph(f"Error: Failed to parse content from AI for {filename}.")
             doc.save(filename)
             return False


        # Process the parts
        i = 0
        while i < len(parts):
            part = parts[i]
            content = ""
            element_type = None

            # Check if the current part is a delimiter
            for type_code, delim_str in delimiters.items():
                if part == delim_str:
                    element_type = type_code
                    # The next part should be the content
                    if i + 1 < len(parts):
                        # Check if the content part is NOT another delimiter
                        is_next_part_delimiter = False
                        for next_delim in all_delimiters_flat:
                            if parts[i+1] == next_delim:
                                is_next_part_delimiter = True
                                break
                        if not is_next_part_delimiter:
                             content = parts[i+1]
                             i += 1 # Consume the content part
                        else:
                             # Handle case like @@H1@@@@H2@@ (empty H1 content)
                             content = "" # Assign empty content
                    else:
                        # Handle case like @@H1@@ at the very end (empty content)
                        content = ""
                    break # Found the delimiter type

            # Add content to doc based on the detected element type
            if element_type == "H1":
                doc.add_heading(content, level=1)
            elif element_type == "H2":
                doc.add_heading(content, level=2)
            elif element_type == "P":
                # You can apply specific styles if defined in a template
                # e.g., doc.add_paragraph(content, style='Body Text')
                doc.add_paragraph(content)
            # Add elif for 'L' (list) if implemented:
            # elif element_type == "L":
            #    doc.add_paragraph(content, style='List Bullet')
            elif element_type is None and part not in all_delimiters_flat:
                # If a part is not a delimiter and wasn't consumed as content,
                # treat it as a default paragraph (might happen with malformed input)
                print(f"Warning: Treating unexpected part as paragraph in {filename}: '{part[:50]}...'")
                doc.add_paragraph(part)

            i += 1 # Move to the next part


        # Save the document
        doc.save(filename)
        print(f"Successfully saved structured report to {filename}")
        return True

    except Exception as e:
        print(f"Error generating structured .docx file '{filename}': {e}")
        # Consider logging the raw report_text here for debugging
        # print(f"Raw report text was: {report_text}")
        return False


def print_docx(filename):
    """Attempts to send the specified .docx file to the default printer."""
    if not os.path.exists(filename):
        print(f"Error: Cannot print file '{filename}' because it does not exist.")
        return

    platform = sys.platform
    try:
        if platform.startswith('win'):
            print(f"Sending '{filename}' to the printer on Windows...")
            os.startfile(filename, "print")
            print(f"Print command sent for '{filename}' on Windows.")
        elif platform.startswith('darwin') or platform.startswith('linux'):
            print(f"Sending '{filename}' to the printer via 'lp' command...")
            subprocess.run(["lp", filename], check=True, capture_output=True) # Use capture_output to hide lp messages unless error
            print(f"Print command sent for '{filename}' via 'lp'.")
        else:
            print(f"Printing not supported on this operating system: {platform}")
    except FileNotFoundError:
         # Specific error if 'lp' isn't found on Unix-like systems
         print(f"Error printing '{filename}': 'lp' command not found. Ensure CUPS is installed.")
    except subprocess.CalledProcessError as e:
         # Error if 'lp' command fails (e.g., printer issue)
         print(f"Error executing 'lp' command for '{filename}': {e}")
         print(f"lp stderr: {e.stderr.decode()}")
    except Exception as e:
        # Catch any other unexpected errors during printing
        print(f"Unexpected error printing '{filename}' on {platform}: {e}")


# ==============================================================================
# Main Execution Block
# ==============================================================================

def main():
    """
    Main function supporting CLI args for provider/model selection,
    falling back to config.yaml, and saving output to dated subdirectories.
    """
    # --- Argument Parsing ---
    parser = argparse.ArgumentParser(description="Generate news reports using a selected AI provider.")
    parser.add_argument(
        "--provider",
        choices=['openai', 'anthropic', 'google'],
        help="Specify the AI provider to use (overrides config.yaml)."
    )
    parser.add_argument(
        "--model",
        help="Specify the exact model name to use (overrides config.yaml)."
    )
    args = parser.parse_args()

    # --- Initialization ---
    print("Initializing...")
    config = load_config() # Load config for defaults and model lookups
    api_keys = load_api_keys()

    # Determine provider: CLI arg > config file
    selected_provider = args.provider if args.provider else config.get('ai_provider')
    if not selected_provider:
        print("Error: AI provider must be specified either via --provider argument or in config.yaml.")
        sys.exit(1)
    if selected_provider not in ['openai', 'anthropic', 'google']:
         print(f"Error: Invalid provider '{selected_provider}'. Choices are 'openai', 'anthropic', 'google'.")
         sys.exit(1)

    # Determine model: CLI arg > config file lookup for the selected provider
    model_name = args.model # Use CLI model if provided
    if not model_name:
        model_name = config.get('model', {}).get(selected_provider)
        if not model_name:
            print(f"Error: Model for provider '{selected_provider}' must be specified either via --model argument or in config.yaml.")
            sys.exit(1)

    print(f"Using AI Provider: {selected_provider.upper()}, Model: {model_name}")

    # Initialize clients
    ai_client = initialize_ai_client(selected_provider, api_keys.get(selected_provider))
    tavily_client = initialize_tavily_client(api_keys.get("tavily"))

    # --- Define Queries ---
    # queries = {
    #     "local": "Reston, Virginia news",
    #     "regional": "Northern Virginia news",
    #     "national": "United States news",
    #     "global": "Latest global news"
    # }
    queries = config.get('queries') 

    # --- Get Current Date for Output Directory ---
    date_str = datetime.now().strftime("%Y-%m-%d")

    # --- Processing Loop ---
    print("\nStarting news report generation process...")
    generated_files = []

    for level, query in queries.items(): # 'level' is the category (local, regional, etc.)
        print("-" * 30)
        print(f"Processing: {level.title()} News (Query: '{query}')")

        # 1. Search
        print(f"Searching for {level} news...")
        results = search_news(tavily_client, query)

        # 2. Format
        if not results:
            print(f"No news articles found for {level} news.")
            continue
        else:
            print(f"Found {len(results)} results. Generating report with {selected_provider.upper()}...")
            report = format_news_with_ai(ai_client, selected_provider, model_name, results, level)
            if "Report generation failed" in report:
                print(f"Skipping document generation for {level} due to report failure.")
                continue

        # --- Define Output Path and Create Directories ---
        output_base_dir = os.path.join(date_str, selected_provider, level)
        try:
            os.makedirs(output_base_dir, exist_ok=True) # Create dirs if they don't exist
        except OSError as e:
            print(f"Error creating directory '{output_base_dir}': {e}")
            continue # Skip generating file if directory creation fails

        # Construct filename within the subdirectory
        base_filename = f"{level}_news_{selected_provider}.docx"
        filename = os.path.join(output_base_dir, base_filename)

        # 3. Generate DOCX
        print(f"Generating DOCX file: {filename}...")
        if generate_docx(report, filename):
             generated_files.append(filename) # Append full path
        else:
            print(f"Warning: File '{filename}' was not created successfully.")

    # --- Printing Phase ---
    print("=" * 30)
    print("News report generation phase completed.")
    print("=" * 30)

    if not generated_files:
        print("No reports were successfully generated to print.")
    else:
        print(f"\nAttempting to print {len(generated_files)} generated reports...")
        for filename in generated_files: # filename now includes the full path
            print("-" * 30)
            print_docx(filename)

    print("=" * 30)
    print("Script execution finished.")


# Standard Python entry point check
if __name__ == "__main__":
    main()
